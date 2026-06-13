from __future__ import annotations

import json
import os
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from PIL import Image as PILImage


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT_DOCX = DOCS_DIR / "RefLab_Registro_DNDA.docx"
OUTPUT_PDF = DOCS_DIR / "RefLab_Registro_DNDA.pdf"

ACCENT = RGBColor(111, 193, 31)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(86, 99, 115)
LIGHT_FILL = "F4F7F2"
SOFT_BORDER = "D7E4D2"
CODE_FILL = "F6F8FA"

EXCLUDED_DIRS = {
    ".git",
    ".next",
    ".npm-cache",
    "node_modules",
    "__pycache__",
}

EXCLUDED_FILE_PATTERNS = (
    ".env",
    ".log",
    ".tsbuildinfo",
    "package-lock.json",
)

SECTION_TITLES = [
    "1. Portada y datos identificatorios",
    "2. Descripcion general",
    "3. Objetivo del software",
    "4. Modulos principales",
    "5. Arquitectura tecnica",
    "6. Estructura del proyecto",
    "7. Capturas de pantalla",
    "8. Fragmentos de codigo representativos",
    "9. Cierre",
]

SCREENSHOT_TITLES = [
    "Pantalla principal",
    "Dashboard",
    "Entrenamiento",
    "Evaluaciones",
    "VAR Lab",
    "Rendimiento",
    "Biblioteca",
    "Perfil / RefCard",
]

SCREENSHOT_LABELS = {
    "01_home_pantalla_principal": "Pantalla principal / Home",
    "02_home_modulos_bloqueados": "Home - accesos y modulos bloqueados",
    "03_dashboard_tecnico": "Dashboard tecnico",
    "04_entrenamiento": "Entrenamiento",
    "05_evaluaciones": "Evaluaciones",
    "06_ref_performance_readiness": "Ref Performance - readiness",
    "07_ref_performance_metricas": "Ref Performance - metricas",
    "08_rendimiento_modulos": "Rendimiento - modulos de analisis",
    "09_biblioteca": "Biblioteca IFAB",
    "10_institucional": "Acceso institucional",
    "11_perfil_refcard": "Perfil / RefCard",
    "12_perfil_arbitral": "Perfil arbitral",
    "13_notificaciones": "Notificaciones inteligentes",
    "14_soporte": "Centro de soporte",
    "15_admin": "Panel Admin",
    "16_identidad_visual_reflab": "Identidad visual RefLab",
}

PDF_TEXT_REPLACEMENTS = {
    "Anio": "Año",
    "Indice": "Índice",
    "Descripcion": "Descripción",
    "descripcion": "descripción",
    "Direccion": "Dirección",
    "direccion": "dirección",
    "computacion": "computación",
    "aplicacion": "aplicación",
    "evaluacion": "evaluación",
    "evaluaciones": "evaluaciones",
    "futbol": "fútbol",
    "simulacion": "simulación",
    "analisis": "análisis",
    "Analisis": "Análisis",
    "capacitacion": "capacitación",
    "desempeno": "desempeño",
    "estadistico": "estadístico",
    "tecnico": "técnico",
    "tecnica": "técnica",
    "arbitro": "árbitro",
    "arbitral": "arbitral",
    "practica": "práctica",
    "metodica": "metódica",
    "medicion": "medición",
    "evolucion": "evolución",
    "deposito": "depósito",
    "codigo": "código",
    "identificacion": "identificación",
    "exposicion": "exposición",
    "Autenticacion": "Autenticación",
    "integracion": "integración",
    "Metricas": "Métricas",
    "metricas": "métricas",
    "examenes": "exámenes",
    "topico": "tópico",
    "Topico": "Tópico",
    "Modulo": "Módulo",
    "Modulos": "Módulos",
    "modulo": "módulo",
    "modulos": "módulos",
    "publica": "pública",
    "publico": "público",
    "academicos": "académicos",
    "Academicos": "Académicos",
    "generacion": "generación",
    "ejecucion": "ejecución",
    "persistencia": "persistencia",
    "seleccionados": "seleccionados",
    "presentacion": "presentación",
    "version": "versión",
    "incorporar": "incorporar",
    "credito": "crédito",
}

MODULES = [
    ("Dashboard", "Panel de seguimiento del usuario con resumen de actividad, metricas y recomendaciones de entrenamiento."),
    ("Entrenamiento", "Modulo de practica con clips y ejercicios por topico arbitral, dificultad y criterio tecnico."),
    ("Evaluaciones", "Instancias de examen y medicion de conocimiento reglamentario y toma de decisiones."),
    ("VAR Lab", "Entrenamiento especifico sobre revision VAR, OFR, error claro y manifiesto y decisiones finales."),
    ("Ingles arbitral", "Area de comunicacion arbitral e ingles IFAB para explicar decisiones y practicar vocabulario tecnico."),
    ("Rendimiento", "Sistema de indicadores, historial, radar tecnico, evolucion y analisis por criterio."),
    ("Biblioteca", "Acceso a material reglamentario, IFAB, protocolos, circulares y recursos academicos."),
    ("Perfil", "Gestion de datos personales, identidad RefLab, avatar y preferencias publicas del usuario."),
    ("RefCard", "Credencial digital publica del arbitro con identificador, metricas, QR y datos verificables."),
    ("Ranking", "Vista comparativa de rendimiento y participacion entre usuarios habilitados."),
]

CODE_SNIPPETS = [
    ("app/dashboard/page.tsx", 1, 118, "Ruta de Dashboard y consumo centralizado de metricas"),
    ("components/TrainingClient.tsx", 1, 96, "Componente de entrenamiento y carga de clips"),
    ("components/EnglishExercise.tsx", 1, 100, "Modulo Comunicacion Arbitral / Ingles IFAB"),
    ("lib/performance.ts", 1, 88, "Tipos principales del sistema de rendimiento"),
    ("lib/performance.ts", 150, 212, "Construccion general del dataset de rendimiento"),
    ("lib/refCard.ts", 1, 45, "Identificador publico y URL de RefCard"),
    ("app/api/profile/route.ts", 1, 88, "Ruta de perfil y persistencia general"),
]


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=SOFT_BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(row.cells[idx])


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_paragraph(doc, text="", style=None, size=10.5, color=DARK, bold=False, italic=False, before=0, after=6, align=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.12
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=16 if level == 1 else 13 if level == 2 else 11.5,
        color=ACCENT if level == 1 else DARK,
        bold=True,
    )
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=DARK)


def add_note_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [16.0])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_run_font(run, size=10.5, color=DARK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(body)
    set_run_font(run2, size=9.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [16.0])
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code.splitlines()):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line[:132])
        set_run_font(run, name="Consolas", size=7.2, color=RGBColor(31, 41, 55))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def detect_framework():
    package_path = ROOT / "package.json"
    package = json.loads(package_path.read_text(encoding="utf-8"))
    deps = {**package.get("dependencies", {}), **package.get("devDependencies", {})}
    return {
        "name": package.get("name", "reflab"),
        "next": deps.get("next", "No detectado"),
        "react": deps.get("react", "No detectado"),
        "typescript": deps.get("typescript", "No detectado"),
        "supabase": deps.get("@supabase/supabase-js", "No detectado"),
        "clerk": deps.get("@clerk/nextjs", "No detectado"),
        "vercel": "vercel.json presente" if (ROOT / "vercel.json").exists() else "No detectado",
    }


def safe_tree(max_depth=3):
    lines = ["reflab/"]

    def walk(path, prefix="", depth=0):
        if depth >= max_depth:
            return
        entries = []
        for item in sorted(path.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower())):
            if item.name in EXCLUDED_DIRS:
                continue
            if item.is_file() and item.name.startswith(EXCLUDED_FILE_PATTERNS):
                continue
            if item.is_file() and any(item.name.endswith(pattern) for pattern in EXCLUDED_FILE_PATTERNS):
                continue
            if item.name == "build_reflab_dnda_doc.py":
                continue
            if item.name.endswith((".png", ".jpg", ".jpeg", ".ico", ".svg", ".wav")) and depth > 0:
                continue
            entries.append(item)
        for index, item in enumerate(entries[:42]):
            connector = "`-- " if index == len(entries[:42]) - 1 else "|-- "
            name = item.name + ("/" if item.is_dir() else "")
            lines.append(prefix + connector + name)
            if item.is_dir():
                extension = "    " if index == len(entries[:42]) - 1 else "|   "
                walk(item, prefix + extension, depth + 1)

    walk(ROOT)
    return "\n".join(lines[:165])


def find_screenshots():
    folders = [ROOT / "screenshots", ROOT / "public" / "screenshots"]
    images = []
    for folder in folders:
        if folder.exists():
            for ext in ("*.png", "*.jpg", "*.jpeg"):
                images.extend(sorted(folder.rglob(ext)))
    return images


def screenshot_label(image):
    return SCREENSHOT_LABELS.get(image.stem, image.stem.replace("_", " ").title())


def sanitize_code(text):
    text = re.sub(r"(sk-[A-Za-z0-9_\-]{12,})", "[REDACTED_API_KEY]", text)
    text = re.sub(r"(pk_(?:test|live)_[A-Za-z0-9_\-]{12,})", "[REDACTED_PUBLIC_KEY]", text)
    text = re.sub(r"(eyJ[A-Za-z0-9_\-\.]{24,})", "[REDACTED_TOKEN]", text)
    text = re.sub(r"(?i)(password|secret|service_role|api_key)\s*[:=]\s*['\"][^'\"]+['\"]", r"\1: \"[REDACTED]\"", text)
    return text


def extract_snippet(relative_path, start, end):
    path = ROOT / relative_path
    if not path.exists():
        return ""
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    selected = lines[max(start - 1, 0): min(end, len(lines))]
    numbered = [f"{start + idx:>4} | {line}" for idx, line in enumerate(selected)]
    return sanitize_code("\n".join(numbered))


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(19)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for idx, (size, color) in enumerate(((16, ACCENT), (13, DARK), (11.5, DARK)), start=1):
        style = styles[f"Heading {idx}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(14 if idx == 1 else 9)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run("RefLab | Material identificatorio DNDA | Pagina ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(footer_p)

    return doc


def add_cover(doc):
    logo = ROOT / "public" / "rf-logo.png"
    add_paragraph(doc, "MATERIAL IDENTIFICATORIO DE SOFTWARE", size=9.5, color=ACCENT, bold=True, after=18)
    if logo.exists():
        try:
            doc.add_picture(str(logo), width=Cm(2.0))
        except Exception:
            pass
    add_paragraph(doc, "RefLab", size=38, color=DARK, bold=True, before=24, after=2)
    add_paragraph(doc, "Plataforma de entrenamiento arbitral", size=16, color=MUTED, after=28)

    table = doc.add_table(rows=5, cols=2)
    set_table_width(table, [4.2, 11.8])
    rows = [
        ("Autor", "David Ricardo Cornejo"),
        ("Anio", "2026"),
        ("Tipo de obra", "Programa de computacion"),
        ("Destino", "Direccion Nacional del Derecho de Autor"),
        ("Documento", "Material identificatorio sin exposicion de codigo fuente completo ni secretos comerciales"),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = table.rows[idx].cells
        shade_cell(cells[0], LIGHT_FILL)
        for cell in cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_run_font(cells[0].paragraphs[0].add_run(label), size=9.5, color=DARK, bold=True)
        set_run_font(cells[1].paragraphs[0].add_run(value), size=9.5, color=DARK)

    add_paragraph(
        doc,
        "Este documento se prepara como soporte descriptivo y tecnico de la obra de software RefLab. "
        "No contiene credenciales, claves API, tokens, variables privadas ni codigo fuente completo.",
        size=10.5,
        color=MUTED,
        before=24,
        after=10,
    )
    add_paragraph(doc, f"Fecha de generacion: {date.today().isoformat()}", size=9.5, color=MUTED, after=0)
    doc.add_page_break()


def add_index(doc):
    add_heading(doc, "Indice", level=1)
    add_paragraph(
        doc,
        "El presente indice organiza el material identificatorio de la plataforma RefLab como programa de computacion.",
        color=MUTED,
    )
    for title in SECTION_TITLES[1:]:
        add_paragraph(doc, title, size=10.5, color=DARK, bold=True, after=2)
    doc.add_page_break()


def add_description(doc):
    add_heading(doc, "2. Descripcion general", level=1)
    add_paragraph(
        doc,
        "RefLab es una plataforma digital destinada al entrenamiento, evaluacion y desarrollo de arbitros de futbol "
        "mediante simulacion de jugadas, analisis de decisiones, capacitacion reglamentaria, evaluacion de desempeno "
        "y seguimiento estadistico.",
    )
    add_paragraph(
        doc,
        "La aplicacion combina interfaces de entrenamiento, ejercicios con clips, evaluaciones, dashboard tecnico, "
        "perfiles de usuario, credenciales digitales y herramientas institucionales. Su finalidad es asistir al arbitro "
        "en la practica metodica de decisiones de partido y en la medicion objetiva de su evolucion.",
    )
    add_note_box(
        doc,
        "Alcance del documento",
        "El material identifica componentes, arquitectura y fragmentos representativos. No sustituye al deposito de codigo fuente ni revela secretos comerciales.",
    )


def add_objective(doc):
    add_heading(doc, "3. Objetivo del software", level=1)
    add_paragraph(doc, "RefLab permite al usuario arbitral e institucional:")
    for item in [
        "Entrenar la toma de decisiones arbitrales con escenarios de juego.",
        "Analizar jugadas y justificar criterios tecnicos.",
        "Evaluar conocimientos reglamentarios y protocolos vinculados al arbitraje.",
        "Medir rendimiento por modulo, topico y criterio tecnico.",
        "Registrar evolucion individual del usuario mediante intentos, examenes, metricas y credenciales.",
    ]:
        add_bullet(doc, item)


def add_modules(doc):
    add_heading(doc, "4. Modulos principales", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [4.2, 11.8])
    table.rows[0].cells[0].text = "Modulo"
    table.rows[0].cells[1].text = "Funcion identificatoria"
    for cell in table.rows[0].cells:
        shade_cell(cell, LIGHT_FILL)
        cell.paragraphs[0].runs[0].bold = True
    for name, detail in MODULES:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = detail
    set_table_width(table, [4.2, 11.8])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, color=DARK, bold=run.bold)


def add_architecture(doc, framework):
    add_heading(doc, "5. Arquitectura tecnica", level=1)
    add_paragraph(
        doc,
        "La arquitectura combina una aplicacion web moderna con servicios gestionados para autenticacion, datos, almacenamiento "
        "y despliegue. La descripcion siguiente es general y no incluye claves ni configuraciones sensibles.",
    )
    rows = [
        ("Frontend", f"Next.js {framework['next']} / React {framework['react']} / TypeScript {framework['typescript']}"),
        ("Backend / datos", f"Supabase y base de datos relacional PostgreSQL; cliente detectado: {framework['supabase']}"),
        ("Autenticacion", f"Clerk; integracion detectada: {framework['clerk']}"),
        ("Hosting", f"Vercel; {framework['vercel']}"),
        ("UI", "Componentes reutilizables, AppShell, modulos de entrenamiento, formularios y vistas de rendimiento."),
        ("Metricas", "Servicio centralizado de rendimiento para intentos, examenes, topicos, criterios, ranking y RefCard."),
        ("Persistencia", "Tablas de perfiles, roles, intentos, examenes, clips, notificaciones y datos institucionales."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [4.3, 11.7])
    table.rows[0].cells[0].text = "Capa"
    table.rows[0].cells[1].text = "Descripcion"
    for cell in table.rows[0].cells:
        shade_cell(cell, LIGHT_FILL)
        cell.paragraphs[0].runs[0].bold = True
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_table_width(table, [4.3, 11.7])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, color=DARK, bold=run.bold)


def add_project_structure(doc):
    add_heading(doc, "6. Estructura del proyecto", level=1)
    add_paragraph(
        doc,
        "Arbol resumido generado automaticamente. Se excluyen .env, node_modules, .next, claves privadas, tokens, logs y archivos sensibles.",
        color=MUTED,
    )
    add_code_block(doc, safe_tree())


def add_screenshots(doc):
    add_heading(doc, "7. Capturas de pantalla", level=1)
    screenshots = find_screenshots()
    if screenshots:
        add_paragraph(doc, "Se detectaron capturas locales y se insertan como referencia visual.", color=MUTED)
        for image in screenshots:
            add_paragraph(doc, screenshot_label(image), size=10.5, bold=True)
            try:
                doc.add_picture(str(image), width=Cm(15.0))
            except Exception:
                add_placeholder(doc, screenshot_label(image))
    else:
        add_paragraph(
            doc,
            "No se detecto carpeta /screenshots ni /public/screenshots. Se dejan placeholders para incorporar capturas al presentar la version final.",
            color=MUTED,
        )
        table = doc.add_table(rows=4, cols=2)
        set_table_width(table, [8.0, 8.0])
        for idx, title in enumerate(SCREENSHOT_TITLES):
            cell = table.cell(idx // 2, idx % 2)
            shade_cell(cell, "F8FAFC")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(18)
            paragraph.paragraph_format.space_after = Pt(18)
            run = paragraph.add_run(f"[Placeholder]\n{title}")
            set_run_font(run, size=10, color=MUTED, bold=True)


def add_placeholder(doc, title):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [16.0])
    cell = table.cell(0, 0)
    shade_cell(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"[Placeholder] {title}")
    set_run_font(run, size=10, color=MUTED, bold=True)


def add_code_snippets(doc):
    add_heading(doc, "8. Fragmentos de codigo representativos", level=1)
    add_paragraph(
        doc,
        "Los siguientes fragmentos son parciales y fueron seleccionados para identificar estructura, componentes y flujo general. "
        "No incluyen credenciales, variables privadas ni secretos comerciales.",
        color=MUTED,
    )
    for relative_path, start, end, title in CODE_SNIPPETS:
        snippet = extract_snippet(relative_path, start, end)
        if not snippet:
            continue
        add_heading(doc, f"{title}", level=2)
        add_paragraph(doc, relative_path, size=8.5, color=MUTED, bold=True, after=3)
        add_code_block(doc, snippet)


def add_closing(doc):
    add_heading(doc, "9. Cierre", level=1)
    add_paragraph(
        doc,
        "El presente documento se acompana como material identificatorio de la obra de software RefLab, a efectos de su registro "
        "como programa de computacion ante la Direccion Nacional del Derecho de Autor.",
    )
    add_paragraph(
        doc,
        "Se deja constancia de que el documento fue preparado con criterio de minimizacion de informacion sensible: no incluye claves API, "
        "tokens, credenciales, secretos de servidor ni codigo fuente completo.",
        color=MUTED,
    )


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "RefLabTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=36,
            leading=42,
            textColor=colors.HexColor("#111827"),
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "subtitle": ParagraphStyle(
            "RefLabSubtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=15,
            leading=20,
            textColor=colors.HexColor("#566373"),
            spaceAfter=22,
        ),
        "kicker": ParagraphStyle(
            "RefLabKicker",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#6FC11F"),
            uppercase=True,
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "RefLabH1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#3F8F16"),
            spaceBefore=12,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "RefLabH2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12.5,
            leading=16,
            textColor=colors.HexColor("#111827"),
            spaceBefore=8,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "RefLabBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.6,
            leading=13,
            textColor=colors.HexColor("#111827"),
            spaceAfter=6,
        ),
        "muted": ParagraphStyle(
            "RefLabMuted",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=12,
            textColor=colors.HexColor("#566373"),
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "RefLabBullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.4,
            leading=12.5,
            leftIndent=14,
            firstLineIndent=-8,
            textColor=colors.HexColor("#111827"),
            spaceAfter=3,
        ),
        "small_bold": ParagraphStyle(
            "RefLabSmallBold",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor("#111827"),
            spaceAfter=3,
        ),
        "code": ParagraphStyle(
            "RefLabCode",
            parent=base["Code"],
            fontName="Courier",
            fontSize=6.65,
            leading=8.1,
            textColor=colors.HexColor("#1F2937"),
            backColor=colors.HexColor("#F6F8FA"),
            borderColor=colors.HexColor("#D7E4D2"),
            borderWidth=0.35,
            borderPadding=5,
            spaceBefore=2,
            spaceAfter=7,
        ),
        "footer": ParagraphStyle(
            "RefLabFooter",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=7.5,
            textColor=colors.HexColor("#566373"),
            alignment=TA_RIGHT,
        ),
        "center": ParagraphStyle(
            "RefLabCenter",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#566373"),
            alignment=TA_CENTER,
        ),
    }


def pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#566373"))
    canvas.drawRightString(
        A4[0] - 1.8 * cm,
        1.1 * cm,
        f"RefLab | Material identificatorio DNDA | Pagina {doc.page}",
    )
    canvas.restoreState()


def pdf_table(data, widths, header=True):
    table = Table(data, colWidths=widths, hAlign="CENTER", repeatRows=1 if header else 0)
    style = [
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D7E4D2")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    if header:
        style.extend([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F7F2")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111827")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ])
    table.setStyle(TableStyle(style))
    return table


def pdf_image(image, max_width=16.2 * cm, max_height=8.1 * cm):
    with PILImage.open(image) as opened:
        width, height = opened.size
    ratio = min(max_width / width, max_height / height)
    return PdfImage(str(image), width=width * ratio, height=height * ratio)


def p(text, style):
    for source, target in PDF_TEXT_REPLACEMENTS.items():
        text = text.replace(source, target)
    return Paragraph(text.replace("&", "&amp;"), style)


def code_chunks(code, max_lines=38):
    lines = [line[:126] for line in code.splitlines()]
    for index in range(0, len(lines), max_lines):
        yield "\n".join(lines[index:index + max_lines])


def add_pdf_code(story, styles, code):
    for chunk in code_chunks(code):
        story.append(Preformatted(chunk, styles["code"]))


def build_pdf():
    styles = pdf_styles()
    framework = detect_framework()
    story = []
    logo = ROOT / "public" / "rf-logo.png"

    story.append(p("MATERIAL IDENTIFICATORIO DE SOFTWARE", styles["kicker"]))
    if logo.exists():
        try:
            story.append(PdfImage(str(logo), width=1.8 * cm, height=1.8 * cm))
            story.append(Spacer(1, 0.55 * cm))
        except Exception:
            pass
    story.append(p("RefLab", styles["title"]))
    story.append(p("Plataforma de entrenamiento arbitral", styles["subtitle"]))
    cover_data = [
        [p("<b>Autor</b>", styles["body"]), p("David Ricardo Cornejo", styles["body"])],
        [p("<b>Anio</b>", styles["body"]), p("2026", styles["body"])],
        [p("<b>Tipo de obra</b>", styles["body"]), p("Programa de computacion", styles["body"])],
        [p("<b>Destino</b>", styles["body"]), p("Direccion Nacional del Derecho de Autor", styles["body"])],
        [p("<b>Documento</b>", styles["body"]), p("Material identificatorio sin exposicion de codigo fuente completo ni secretos comerciales", styles["body"])],
    ]
    story.append(pdf_table(cover_data, [4.0 * cm, 12.2 * cm], header=False))
    story.append(Spacer(1, 0.65 * cm))
    story.append(p(
        "Este documento se prepara como soporte descriptivo y tecnico de la obra de software RefLab. "
        "No contiene credenciales, claves API, tokens, variables privadas ni codigo fuente completo.",
        styles["muted"],
    ))
    story.append(p(f"Fecha de generacion: {date.today().isoformat()}", styles["muted"]))
    story.append(PageBreak())

    story.append(p("Indice", styles["h1"]))
    story.append(p("El presente indice organiza el material identificatorio de la plataforma RefLab como programa de computacion.", styles["muted"]))
    for title in SECTION_TITLES[1:]:
        story.append(p(f"<b>{title}</b>", styles["body"]))
    story.append(PageBreak())

    story.append(p("2. Descripcion general", styles["h1"]))
    story.append(p(
        "RefLab es una plataforma digital destinada al entrenamiento, evaluacion y desarrollo de arbitros de futbol "
        "mediante simulacion de jugadas, analisis de decisiones, capacitacion reglamentaria, evaluacion de desempeno "
        "y seguimiento estadistico.",
        styles["body"],
    ))
    story.append(p(
        "La aplicacion combina interfaces de entrenamiento, ejercicios con clips, evaluaciones, dashboard tecnico, "
        "perfiles de usuario, credenciales digitales y herramientas institucionales. Su finalidad es asistir al arbitro "
        "en la practica metodica de decisiones de partido y en la medicion objetiva de su evolucion.",
        styles["body"],
    ))
    story.append(pdf_table(
        [[p("<b>Alcance del documento</b><br/>El material identifica componentes, arquitectura y fragmentos representativos. No sustituye al deposito de codigo fuente ni revela secretos comerciales.", styles["body"])]],
        [16.2 * cm],
        header=False,
    ))
    story.append(Spacer(1, 0.18 * cm))

    story.append(p("3. Objetivo del software", styles["h1"]))
    story.append(p("RefLab permite al usuario arbitral e institucional:", styles["body"]))
    for item in [
        "Entrenar la toma de decisiones arbitrales con escenarios de juego.",
        "Analizar jugadas y justificar criterios tecnicos.",
        "Evaluar conocimientos reglamentarios y protocolos vinculados al arbitraje.",
        "Medir rendimiento por modulo, topico y criterio tecnico.",
        "Registrar evolucion individual del usuario mediante intentos, examenes, metricas y credenciales.",
    ]:
        story.append(p(f"- {item}", styles["bullet"]))

    story.append(p("4. Modulos principales", styles["h1"]))
    module_rows = [[p("<b>Modulo</b>", styles["body"]), p("<b>Funcion identificatoria</b>", styles["body"])]]
    for name, detail in MODULES:
        module_rows.append([p(f"<b>{name}</b>", styles["body"]), p(detail, styles["body"])])
    story.append(pdf_table(module_rows, [4.1 * cm, 12.1 * cm]))
    story.append(Spacer(1, 0.15 * cm))

    story.append(p("5. Arquitectura tecnica", styles["h1"]))
    story.append(p(
        "La arquitectura combina una aplicacion web moderna con servicios gestionados para autenticacion, datos, almacenamiento "
        "y despliegue. La descripcion siguiente es general y no incluye claves ni configuraciones sensibles.",
        styles["body"],
    ))
    arch_rows = [[p("<b>Capa</b>", styles["body"]), p("<b>Descripcion</b>", styles["body"])]]
    for label, value in [
        ("Frontend", f"Next.js {framework['next']} / React {framework['react']} / TypeScript {framework['typescript']}"),
        ("Backend / datos", f"Supabase y base de datos relacional PostgreSQL; cliente detectado: {framework['supabase']}"),
        ("Autenticacion", f"Clerk; integracion detectada: {framework['clerk']}"),
        ("Hosting", f"Vercel; {framework['vercel']}"),
        ("UI", "Componentes reutilizables, AppShell, modulos de entrenamiento, formularios y vistas de rendimiento."),
        ("Metricas", "Servicio centralizado de rendimiento para intentos, examenes, topicos, criterios, ranking y RefCard."),
        ("Persistencia", "Tablas de perfiles, roles, intentos, examenes, clips, notificaciones y datos institucionales."),
    ]:
        arch_rows.append([p(f"<b>{label}</b>", styles["body"]), p(value, styles["body"])])
    story.append(pdf_table(arch_rows, [4.2 * cm, 12.0 * cm]))
    story.append(PageBreak())

    story.append(p("6. Estructura del proyecto", styles["h1"]))
    story.append(p(
        "Arbol resumido generado automaticamente. Se excluyen .env, node_modules, .next, claves privadas, tokens, logs y archivos sensibles.",
        styles["muted"],
    ))
    add_pdf_code(story, styles, safe_tree())

    story.append(p("7. Capturas de pantalla", styles["h1"]))
    screenshots = find_screenshots()
    if screenshots:
        story.append(p("Se detectaron capturas locales y se insertan como referencia visual.", styles["muted"]))
        for image in screenshots:
            try:
                story.append(KeepTogether([
                    p(f"<b>{screenshot_label(image)}</b>", styles["body"]),
                    pdf_image(image),
                    Spacer(1, 0.2 * cm),
                ]))
            except Exception:
                story.append(pdf_table([[p(f"[Placeholder] {screenshot_label(image)}", styles["center"])]], [16.2 * cm], header=False))
    else:
        story.append(p(
            "No se detecto carpeta /screenshots ni /public/screenshots. Se dejan placeholders para incorporar capturas al presentar la version final.",
            styles["muted"],
        ))
        placeholder_rows = []
        for index in range(0, len(SCREENSHOT_TITLES), 2):
            placeholder_rows.append([
                p(f"[Placeholder]<br/><b>{SCREENSHOT_TITLES[index]}</b>", styles["center"]),
                p(f"[Placeholder]<br/><b>{SCREENSHOT_TITLES[index + 1]}</b>", styles["center"]),
            ])
        table = pdf_table(placeholder_rows, [8.1 * cm, 8.1 * cm], header=False)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
            ("ROWHEIGHT", (0, 0), (-1, -1), 1.9 * cm),
        ]))
        story.append(table)
    story.append(PageBreak())

    story.append(p("8. Fragmentos de codigo representativos", styles["h1"]))
    story.append(p(
        "Los siguientes fragmentos son parciales y fueron seleccionados para identificar estructura, componentes y flujo general. "
        "No incluyen credenciales, variables privadas ni secretos comerciales.",
        styles["muted"],
    ))
    for relative_path, start, end, title in CODE_SNIPPETS:
        snippet = extract_snippet(relative_path, start, end)
        if not snippet:
            continue
        story.append(p(title, styles["h2"]))
        story.append(p(relative_path, styles["small_bold"]))
        add_pdf_code(story, styles, snippet)

    story.append(p("9. Cierre", styles["h1"]))
    story.append(p(
        "El presente documento se acompana como material identificatorio de la obra de software RefLab, a efectos de su registro "
        "como programa de computacion ante la Direccion Nacional del Derecho de Autor.",
        styles["body"],
    ))
    story.append(p(
        "Se deja constancia de que el documento fue preparado con criterio de minimizacion de informacion sensible: no incluye claves API, "
        "tokens, credenciales, secretos de servidor ni codigo fuente completo.",
        styles["muted"],
    ))

    pdf = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.7 * cm,
        title="RefLab - Registro DNDA",
        author="David Ricardo Cornejo",
    )
    pdf.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)
    print(OUTPUT_PDF)


def main():
    DOCS_DIR.mkdir(exist_ok=True)
    framework = detect_framework()
    doc = setup_document()
    doc.core_properties.title = "RefLab - Registro DNDA"
    doc.core_properties.author = "David Ricardo Cornejo"
    doc.core_properties.subject = "Material identificatorio de programa de computacion"
    doc.core_properties.keywords = "RefLab, programa de computacion, DNDA, arbitraje, software"

    add_cover(doc)
    add_index(doc)
    add_description(doc)
    add_objective(doc)
    add_modules(doc)
    add_architecture(doc, framework)
    add_project_structure(doc)
    add_screenshots(doc)
    add_code_snippets(doc)
    add_closing(doc)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    build_pdf()


if __name__ == "__main__":
    main()
